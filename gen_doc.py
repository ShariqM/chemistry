"""
Create a .DOC file for a list of atoms, one per page.

Each atom has a description and an image of the electron shells of the atom.

Before running this code:
- Images should be in outputs/*.png
- Atom descriptions should be in inputs/atom_descriptions.txt
- Create outputs_doc/ directory, .DOC will be written there.

Dependencies:
- pip install python-docx

"""

from docx import Document
from docx.shared import Inches
import pdb
import glob

def main():
    document = Document()

    # 1/2" margins
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)


    # Sort images by number instead of linux default.
    image_path = "outputs"
    image_filepaths = sorted(glob.glob(f"{image_path}/*.png"),
                             key=lambda x: int(x[x.index('/')+1:x.index('.')]))

    atom_index = 0
    f = open("inputs/atom_descriptions.txt", 'r')
    line = f.readline()
    while line != "":
        # Read description lines until newline or end of file
        lines = []
        while line != '\n' and line != "":
            lines.append(line)
            line = f.readline()
        description = "".join(lines)
        print (description)

        # Add description, image, and a new page to the doc.
        document.add_paragraph(description)
        document.add_picture(image_filepaths[atom_index], width=Inches(6.40))
        document.add_page_break()

        # Move to start of next atom description
        line = f.readline()
        atom_index = atom_index + 1

    # Save document
    document.save('output_doc/atom_shells.docx')

if __name__ == "__main__":
    main()
